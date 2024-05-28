# import json
# from docx import Document

# def parse_docx(file_path):
#     doc = Document(file_path)

#     # Extract class and subject from the header
#     header = doc.sections[0].header
#     class_info = None
#     subject_info = None

#     for paragraph in header.paragraphs:
#         text = paragraph.text.strip()
#         if text.startswith('Class:'):
#             class_info = text.split('Class:')[1].split('Subject:')[0].strip()
#             subject_info = text.split('Subject:')[1].strip()
#             break

#     questions = []
#     current_question = None
#     current_options = {}
#     option_keys = ['a', 'b', 'c', 'd']
#     option_index = 0
#     is_question = False

#     for paragraph in doc.paragraphs:
#         text = paragraph.text.strip()

#         if text.startswith('Answer:'):
#             answer = text.split('Answer:')[1].strip()
#             if current_question:
#                 current_question['options'] = current_options
#                 current_question['correct_answer'] = answer
#                 questions.append(current_question)
#                 current_question = None
#                 current_options = {}
#                 option_index = 0
#             is_question = False
#         elif text and not text.startswith('Class:'):
#             if is_question:
#                 if option_index < len(option_keys):
#                     current_options[option_keys[option_index]] = text
#                     option_index += 1
#             else:
#                 if current_question:
#                     current_question['options'] = current_options
#                     questions.append(current_question)
#                 current_question = {'question': text, 'options': {}}
#                 current_options = {}
#                 option_index = 0
#                 is_question = True

#     # Final question handling if document ends without 'Answer:'
#     if current_question:
#         current_question['options'] = current_options
#         questions.append(current_question)

#     return {
#         'class': class_info,
#         'subject': subject_info,
#         'questions': questions
#     }

# # Path to the DOCX file
# doc_path = 'Sample QP to upload(1).docx'
# data = parse_docx(doc_path)

# # Convert to JSON
# json_data = json.dumps(data, indent=4)

# # Print the JSON data
# print(json_data)

# # Optionally, save to a JSON file
# json_output_path = 'output.json'
# with open(json_output_path, 'w') as json_file:
#     json_file.write(json_data)
from flask import Flask, request, jsonify
from docx import Document
import os

app = Flask(__name__)

def parse_docx(file_path):
    doc = Document(file_path)

    # Extract class and subject from the header
    header = doc.sections[0].header
    class_info = None
    subject_info = None

    for paragraph in header.paragraphs:
        text = paragraph.text.strip()
        if text.startswith('Class:'):
            class_info = text.split('Class:')[1].split('Subject:')[0].strip()
            subject_info = text.split('Subject:')[1].strip()
            break

    questions = []
    current_question = None
    current_options = {}
    option_keys = ['a', 'b', 'c', 'd']
    option_index = 0
    is_question = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text.startswith('Answer:'):
            answer = text.split('Answer:')[1].strip()
            if current_question:
                current_question['options'] = current_options
                current_question['correct_answer'] = answer
                questions.append(current_question)
                current_question = None
                current_options = {}
                option_index = 0
            is_question = False
        elif text and not text.startswith('Class:'):
            if is_question:
                if option_index < len(option_keys):
                    current_options[option_keys[option_index]] = text
                    option_index += 1
            else:
                if current_question:
                    current_question['options'] = current_options
                    questions.append(current_question)
                current_question = {'question': text, 'options': {}}
                current_options = {}
                option_index = 0
                is_question = True

    # Final question handling if document ends without 'Answer:'
    if current_question:
        current_question['options'] = current_options
        questions.append(current_question)

    return {
        'class': class_info,
        'subject': subject_info,
        'questions': questions
    }

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.endswith('.docx'):
        file_path = os.path.join('/tmp', file.filename)
        file.save(file_path)
        data = parse_docx(file_path)
        os.remove(file_path)  # Clean up the file after processing
        return jsonify(data)
    else:
        return jsonify({'error': 'Invalid file type. Only .docx files are allowed.'}), 400

if __name__ == '__main__':
    app.run(debug=True)
